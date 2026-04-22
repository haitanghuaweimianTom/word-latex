"""
LaTeX 到 Word 转换引擎
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from parser import (
    LaTeXDocument, LaTeXElement, TextElement, CommandElement,
    MathElement, EnvironmentElement, TableElement
)
from math_converter import LaTeXToOMMLConverter
from table_converter import LaTeXTableToWordConverter


@dataclass
class ConversionOptions:
    """转换选项"""
    package_mode: bool = False
    verbose: bool = False


class LaTeXToWordConverter:
    """
    LaTeX 到 Word 转换器

    支持转换:
    - 标题 (section, subsection, etc.)
    - 段落文本
    - 数学公式 (转换为 OMML)
    - 表格 (tabular → Word 表格)
    - 列表 (itemize, enumerate)
    - 加粗和斜体
    """

    # Word 内置样式名称
    HEADING_STYLES = {
        1: 'Heading 1',
        2: 'Heading 2',
        3: 'Heading 3',
        4: 'Heading 4',
        5: 'Heading 5',
    }

    def __init__(self, options: ConversionOptions = None):
        self.options = options or ConversionOptions()
        self.math_converter = LaTeXToOMMLConverter()
        self.table_converter = LaTeXTableToWordConverter()

    def convert(self, latex_doc: LaTeXDocument) -> Document:
        """
        将 LaTeX 文档转换为 Word 文档

        Args:
            latex_doc: LaTeXDocument 对象

        Returns:
            python-docx Document 对象
        """
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # 转换每个元素
        for element in latex_doc.elements:
            self._convert_element(doc, element)

        return doc

    def _convert_element(self, doc: Document, element: LaTeXElement):
        """转换单个元素"""
        if isinstance(element, CommandElement):
            self._convert_command(doc, element)
        elif isinstance(element, TextElement):
            self._convert_text(doc, element)
        elif isinstance(element, MathElement):
            self._convert_math(doc, element)
        elif isinstance(element, TableElement):
            self._convert_table(doc, element)
        elif isinstance(element, EnvironmentElement):
            self._convert_environment(doc, element)

    def _convert_command(self, doc: Document, element: CommandElement):
        """转换命令元素"""
        if element.name in ['section', 'subsection', 'subsubsection',
                            'paragraph', 'subparagraph']:
            # 获取标题级别
            levels = {
                'section': 1,
                'subsection': 2,
                'subsubsection': 3,
                'paragraph': 4,
                'subparagraph': 5,
            }
            level = levels.get(element.name, 1)

            # 获取标题文本
            if element.args:
                title = element.args[0]
            else:
                title = ""

            # 添加标题
            heading = doc.add_heading(title, level=level)

            # 应用样式
            style_name = self.HEADING_STYLES.get(level)
            if style_name and style_name in [s.name for s in doc.styles]:
                heading.style = doc.styles[style_name]

    def _convert_text(self, doc: Document, element: TextElement):
        """转换文本元素"""
        para = doc.add_paragraph()

        # 添加文本
        if element.bold and element.italic:
            run = para.add_run(element.content)
            run.bold = True
            run.italic = True
        elif element.bold:
            run = para.add_run(element.content)
            run.bold = True
        elif element.italic:
            run = para.add_run(element.content)
            run.italic = True
        else:
            para.add_run(element.content)

    def _convert_math(self, doc: Document, element: MathElement):
        """转换数学公式"""
        # 添加空行
        if element.display:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            para = doc.add_paragraph()

        # 添加公式（使用纯文本，因为 OMML 过于复杂）
        if element.content:
            math_text = self.math_converter.convert(element.content)
            if math_text:
                para.add_run(math_text)

    def _convert_table(self, doc: Document, element: TableElement):
        """转换表格元素"""
        if not element.rows:
            return

        # 创建表格
        num_rows = len(element.rows)
        num_cols = max(len(row) for row in element.rows)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        # 填充数据
        for i, row_data in enumerate(element.rows):
            row = table.rows[i]

            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = cell_text

                    # 表头加粗
                    if element.has_header and i == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

        # 添加强制段落结束表格
        doc.add_paragraph()

    def _convert_environment(self, doc: Document, element: EnvironmentElement):
        """转换环境元素"""
        if element.name == 'itemize':
            self._convert_unordered_list(doc, element.content)
        elif element.name == 'enumerate':
            self._convert_ordered_list(doc, element.content)

    def _convert_unordered_list(self, doc: Document, items: List[TextElement]):
        """转换无序列表"""
        for item in items:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(item.content)

    def _convert_ordered_list(self, doc: Document, items: List[TextElement]):
        """转换有序列表"""
        for item in items:
            para = doc.add_paragraph(style='List Number')
            para.add_run(item.content)


def convert_latex_to_word(latex_doc: LaTeXDocument) -> Document:
    """便捷函数：将 LaTeX 文档转换为 Word"""
    converter = LaTeXToWordConverter()
    return converter.convert(latex_doc)