"""
测试样例：创建包含各种数学公式的 Word 文档
使用 python-docx 库创建测试文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


def create_formula_test_doc():
    """创建公式测试文档"""
    doc = Document()

    # 设置文档标题
    doc.add_heading('数学公式测试文档', 0)

    # 1. 简单代数公式
    doc.add_heading('1. 基本代数', 1)
    p = doc.add_paragraph()
    run = p.add_run('二次方程: ')
    run.bold = True
    p.add_run('x = (-b ± √(b²-4ac)) / 2a')

    # 2. 分数和指数
    doc.add_heading('2. 分数和指数', 1)
    p = doc.add_paragraph()
    p.add_run('表达式: x² + y² = z²， 其中 (a+b)/c = d')

    # 3. 求和与积分
    doc.add_heading('3. 求和与积分', 1)
    p = doc.add_paragraph()
    p.add_run('求和公式: ∑(i=1 to n) i = n(n+1)/2')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('积分公式: ∫(0 to ∞) e^(-x) dx = 1')

    # 4. 矩阵
    doc.add_heading('4. 矩阵', 1)
    p = doc.add_paragraph()
    p.add_run('2×2 矩阵: [a b; c d]')

    # 5. 希腊字母
    doc.add_heading('5. 希腊字母', 1)
    p = doc.add_paragraph()
    p.add_run('常用希腊字母: α, β, γ, δ, θ, π, σ, ω, Σ, Π, Δ')

    # 保存文档
    doc.save('tests/formula_test.docx')
    print("已创建: tests/formula_test.docx")


def create_table_test_doc():
    """创建表格测试文档"""
    doc = Document()

    doc.add_heading('表格测试', 0)

    # 添加一个 3x4 的表格
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    # 填充表格数据
    headers = ['姓名', '年龄', '城市', '职业']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        # 加粗表头
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 数据行
    data = [
        ['张三', '28', '北京', '工程师'],
        ['李四', '35', '上海', '设计师'],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    doc.save('tests/table_test.docx')
    print("已创建: tests/table_test.docx")


def create_complex_doc():
    """创建包含多种元素的复杂测试文档"""
    doc = Document()

    # 标题
    doc.add_heading('学术论文示例', 0)
    doc.add_heading('摘要', 1)

    abstract = doc.add_paragraph(
        '本文研究了基于深度学习的图像分类方法。'
        '我们提出了一种新的卷积神经网络架构，'
        '在多个基准数据集上取得了最先进的性能。'
    )

    # 关键词
    p = doc.add_paragraph()
    run = p.add_run('关键词: ')
    run.bold = True
    p.add_run('深度学习, 图像分类, 卷积神经网络, 计算机视觉')

    # 第1节：引言
    doc.add_heading('1. 引言', 1)
    doc.add_paragraph(
        '近年来，深度学习技术在计算机视觉领域取得了显著进展。'
        '卷积神经网络 (CNN) 已成为图像分类的主流方法。'
    )

    # 第2节：方法
    doc.add_heading('2. 方法', 1)
    doc.add_heading('2.1 网络架构', 2)

    doc.add_paragraph(
        '我们提出的网络包含以下主要组件：'
    )

    # 无序列表
    items = [
        '卷积层：用于特征提取',
        '池化层：用于降维',
        '全连接层：用于分类',
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    doc.add_heading('2.2 损失函数', 2)
    p = doc.add_paragraph()
    p.add_run('交叉熵损失: ')
    p.add_run('L = -∑(y log(ŷ)))')

    # 第3节：实验
    doc.add_heading('3. 实验结果', 1)

    # 创建一个简单的表格
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['方法', '准确率', '参数量']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    results = [
        ['ResNet-50', '76.5%', '25.6M'],
        ['VGG-16', '72.8%', '138M'],
        ['我们的方法', '78.9%', '12.3M'],
    ]

    for row_idx, row_data in enumerate(results, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    # 第4节：结论
    doc.add_heading('4. 结论', 1)
    doc.add_paragraph(
        '本文提出了一种高效的图像分类方法。'
        '实验结果表明，我们的方法在准确率和效率上都优于现有方法。'
    )

    # 参考文献
    doc.add_heading('参考文献', 1)
    refs = [
        '[1] He, K. et al. (2016). Deep Residual Learning for Image Recognition.',
        '[2] Simonyan, K. & Zisserman, A. (2014). Very Deep Convolutional Networks.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.add_run(ref)

    doc.save('tests/complex_test.docx')
    print("已创建: tests/complex_test.docx")


def create_math_doc():
    """创建纯数学测试文档"""
    doc = Document()

    doc.add_heading('LaTeX 数学公式参考', 0)

    # 内联公式
    doc.add_heading('内联公式', 1)
    p = doc.add_paragraph()
    p.add_run('二次公式: x = (-b ± √(b²-4ac)) / 2a')
    p = doc.add_paragraph()
    p.add_run('勾股定理: a² + b² = c²')

    # 行间公式
    doc.add_heading('行间公式', 1)
    doc.add_paragraph('欧拉公式: e^(iπ) + 1 = 0')
    doc.add_paragraph('积分: ∫(a to b) f(x) dx = F(b) - F(a)')

    # 矩阵
    doc.add_heading('矩阵', 1)
    doc.add_paragraph('标准形式: [1 2; 3 4] 表示 2x2 矩阵')

    # 复杂表达式
    doc.add_heading('复杂表达式', 1)
    doc.add_paragraph('组合数: C(n,k) = n! / (k!(n-k)!)')
    doc.add_paragraph('求和: ∑(i=1 to n) i² = n(n+1)(2n+1)/6')

    doc.save('tests/math_test.docx')
    print("已创建: tests/math_test.docx")


if __name__ == '__main__':
    # 确保测试目录存在
    os.makedirs('tests', exist_ok=True)

    print("创建测试文档...")
    create_formula_test_doc()
    create_table_test_doc()
    create_complex_doc()
    create_math_doc()

    print("\n所有测试文档创建完成!")